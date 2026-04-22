"""Convert the Markdown docs in this folder to .docx files.

Usage:
    python md_to_docx.py

Keeps headings, paragraphs, bold/italic/code spans, bullet and numbered lists,
tables (GFM pipe tables), fenced code blocks, and inline links.
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))

HEADING_SIZES = {1: 22, 2: 18, 3: 14, 4: 12, 5: 11, 6: 11}
MONO_FONT = "Consolas"
BODY_FONT = "Calibri"


def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_inline(paragraph, text):
    """Render inline markdown (bold, italic, code, links) into a paragraph."""
    # Patterns we care about, ordered so we try multi-char tokens first.
    pattern = re.compile(
        r"(\*\*.+?\*\*)"        # bold **text**
        r"|(`[^`]+`)"            # code `text`
        r"|(\*[^*]+\*)"           # italic *text*
        r"|(\[[^\]]+\]\([^)]+\))" # link [text](url)
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = MONO_FONT
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xB0, 0x00, 0x40)
        elif token.startswith("*") and token.endswith("*"):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        elif token.startswith("["):
            link_m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_m:
                label, _url = link_m.group(1), link_m.group(2)
                r = paragraph.add_run(label)
                r.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)
                r.underline = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_heading(doc, level, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    add_inline(p, text)
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(HEADING_SIZES.get(level, 11))
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)


def add_bullet(doc, text, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * indent)
    add_inline(p, text)


def add_number(doc, text, indent=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * indent)
    add_inline(p, text)


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("\n".join(lines))
    r.font.name = MONO_FONT
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # Light gray shading.
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)


def add_table(doc, rows):
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            para = cell.paragraphs[0]
            add_inline(para, cell_text.strip())
            if i == 0:
                for run in para.runs:
                    run.bold = True
                set_cell_shading(cell, "1F2A44")
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def convert(md_path, docx_path):
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)

    i = 0
    in_code = False
    code_buf = []
    table_buf = []

    def flush_table():
        nonlocal table_buf
        if not table_buf:
            return
        # Filter out the separator row `|---|---|`.
        rows = [r for r in table_buf if not re.match(r"^\s*\|[\s:\-|]+\|\s*$", r)]
        parsed = []
        for row in rows:
            cells = row.strip()
            if cells.startswith("|"): cells = cells[1:]
            if cells.endswith("|"): cells = cells[:-1]
            parsed.append([c.strip() for c in cells.split("|")])
        add_table(doc, parsed)
        table_buf = []

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.startswith("```"):
            if not in_code:
                flush_table()
                in_code = True
                code_buf = []
            else:
                add_code_block(doc, code_buf)
                in_code = False
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Pipe table (simple detection: line with `|`)
        if "|" in line and line.strip().startswith("|") and line.strip().endswith("|"):
            table_buf.append(line)
            i += 1
            continue
        else:
            flush_table()

        # Headings
        m = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if m:
            add_heading(doc, len(m.group(1)), m.group(2))
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^\s*---\s*$", line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Bullet / numbered list
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent = len(m.group(1)) // 2
            add_bullet(doc, m.group(2), indent)
            i += 1
            continue
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            indent = len(m.group(1)) // 2
            add_number(doc, m.group(2), indent)
            i += 1
            continue

        # Blank line → spacer paragraph (skip consecutive blanks)
        if not line.strip():
            i += 1
            continue

        # Regular paragraph (join wrapped lines)
        buf = [line]
        j = i + 1
        while j < len(lines) and lines[j].strip() and not lines[j].startswith(("#", "- ", "* ", "```", "|")) and not re.match(r"^\s*\d+\.\s+", lines[j]):
            buf.append(lines[j])
            j += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(s.strip() for s in buf))
        i = j

    flush_table()
    doc.save(docx_path)
    print(f"wrote {docx_path}")


if __name__ == "__main__":
    for md in ("DEVELOPER.md", "USER_GUIDE.md"):
        src = os.path.join(HERE, md)
        if not os.path.exists(src):
            print(f"skip missing {src}")
            continue
        dst = os.path.join(HERE, md.replace(".md", ".docx"))
        convert(src, dst)
